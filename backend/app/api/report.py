from fastapi import APIRouter
from fastapi.responses import FileResponse
from docx import Document
from datetime import datetime

from app.core.database import SessionLocal
from app.models.event_db import EventDB

router = APIRouter()


@router.get("/generate-report")
def generate_report():

    db = SessionLocal()

    total_events = db.query(EventDB).count()

    zone_intrusions = (
        db.query(EventDB)
        .filter(EventDB.event_type == "zone_intrusion")
        .count()
    )

    loitering = (
        db.query(EventDB)
        .filter(EventDB.event_type == "loitering")
        .count()
    )

    high_risk = (
        db.query(EventDB)
        .filter(EventDB.severity == "L3")
        .count()
    )

    recent_events = (
        db.query(EventDB)
        .order_by(EventDB.id.desc())
        .limit(10)
        .all()
    )

    db.close()

    doc = Document()

    doc.add_heading(
        "AI Surveillance Incident Report",
        level=1
    )

    doc.add_paragraph(
        f"Generated: {datetime.now()}"
    )

    doc.add_heading(
        "Summary",
        level=2
    )

    doc.add_paragraph(
        f"Total Events: {total_events}"
    )

    doc.add_paragraph(
        f"Zone Intrusions: {zone_intrusions}"
    )

    doc.add_paragraph(
        f"Loitering Events: {loitering}"
    )

    doc.add_paragraph(
        f"High Risk Events: {high_risk}"
    )

    doc.add_heading(
        "Recent Alerts",
        level=2
    )

    for event in recent_events:
        doc.add_paragraph(
            f"{event.event_type} | "
            f"{event.severity} | "
            f"{event.timestamp}"
        )

    filename = "incident_report.docx"

    doc.save(filename)

    return FileResponse(
        filename,
        media_type=
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename
    )